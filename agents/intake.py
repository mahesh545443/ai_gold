import os
import re
import docx
import pdfplumber


class IntakeAgent:
    def __init__(self):
        pass

    def read_file(self, file_path: str) -> dict:
        """
        Reads a single uploaded file (txt, docx, pdf) and returns structured raw data.
        """
        filename = os.path.basename(file_path)
        print(f"📂 Intake Agent: Reading '{filename}'...")

        try:
            content = ""
            file_type = ""

            if filename.endswith(".txt"):
                file_type = "email_text"
                content = self._read_text(file_path)

            elif filename.endswith(".docx"):
                file_type = "word_document"
                content = self._read_word(file_path)

            elif filename.endswith(".pdf"):
                file_type = "pdf_document"
                content = self._read_pdf(file_path)

            else:
                return {"error": f"Unsupported file type: {filename}"}

            sender_email = self._extract_email(content)
            sender_company = self._extract_company(content)

            print(f"   ✅ Extracted | Email: {sender_email} | Company: {sender_company}")

            return {
                "filename": filename,
                "file_type": file_type,
                "raw_content": content,
                "sender_email": sender_email,
                "sender_company": sender_company,
                "llm_prompt_context": f"File: {filename}\nContent:\n{content}"
            }

        except Exception as e:
            return {"error": f"Failed to read {filename}: {str(e)}"}

    # ── Helpers ────────────────────────────────────────────────

    def _read_text(self, path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    def _read_word(self, path):
        doc = docx.Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                lines.append("| " + " | ".join(c.text.strip() for c in row.cells) + " |")
        return "\n".join(lines)

    def _read_pdf(self, path):
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                pages.append(page.extract_text(layout=True) or "")
        return "\n".join(pages)

    def _extract_email(self, text):
        matches = re.findall(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
        return matches[0] if matches else None

    def _extract_company(self, text):
        patterns = [
            r"(?:Company|From|Organisation|Organization|Firm)\s*:\s*(.+)",
            r"([A-Z][a-zA-Z\s]+(?:Pvt\.?\s*Ltd\.?|Ltd\.?|Inc\.?|Corp\.?|LLP))"
        ]
        for pat in patterns:
            m = re.search(pat, text)
            if m:
                return m.group(1).strip()
        return "Valued Customer"
